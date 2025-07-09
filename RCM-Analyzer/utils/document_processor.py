import pandas as pd
import os
import json
from PyPDF2 import PdfReader
import pdfplumber
import docx
from typing import Dict, List, Any, Union
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def process_document(file_path: str) -> Dict[str, Any]:
    """
    Process different document types and extract structured data
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Structured data extracted from the document
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.xlsx':
            return process_excel(file_path)
        elif file_ext == '.csv':
            return process_csv(file_path)
        elif file_ext == '.pdf':
            return process_pdf(file_path)
        elif file_ext == '.docx':
            return process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise

def process_excel(file_path: str) -> Dict[str, Any]:
    """Process Excel files and extract RCM data"""
    logger.info(f"Processing Excel file: {file_path}")
    
    # Read all sheets
    excel_file = pd.ExcelFile(file_path)
    sheet_names = excel_file.sheet_names
    
    # Initialize the structured data
    structured_data = {
        "metadata": {
            "file_name": os.path.basename(file_path),
            "file_type": "excel",
            "sheet_count": len(sheet_names)
        },
        "raw_data": [],
        "control_objectives": [],
        "risks": [],
        "controls": [],
        "gaps": [],
        "departments": []
    }
    
    # Known RCM column patterns based on the screenshot
    rcm_column_patterns = {
        'type_of_risk': ['type of risk', 'risk type', 'risk category'],
        'area': ['area', 'department', 'function'],
        'control_number': ['control number', 'control no', 'control id', 'control #'],
        'area_subprocess': ['area/ sub process', 'area/sub process', 'sub process', 'process'],
        'control_objective': ['control objective', 'objective', 'control obj'],
        'risk': ['risk/ what can go wrong', 'what can go wrong', 'risk', 'risk description']
    }
    
    # Process each sheet for raw data first
    for sheet_name in sheet_names:
        logger.info(f"Processing sheet: {sheet_name}")
        
        try:
            # Try to read the sheet first
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Clean up column names and drop empty rows/columns
            df.columns = [str(col).strip() for col in df.columns]
            df = df.dropna(how='all').reset_index(drop=True)
            
            # Log the column names for debugging
            logger.info(f"Columns in sheet {sheet_name}: {list(df.columns)}")
            
            # Store the raw data regardless of format
            sheet_data = {
                "sheet_name": sheet_name,
                "rows": []
            }
            
            for idx, row in df.iterrows():
                # Skip completely empty rows
                if sum(pd.notna(row)) < 1:
                    continue
                
                # Add all row data
                row_data = {col: str(row[col]) if pd.notna(row[col]) else "" for col in df.columns}
                sheet_data["rows"].append(row_data)
            
            structured_data["raw_data"].append(sheet_data)
            logger.info(f"Stored {len(sheet_data['rows'])} raw rows from sheet {sheet_name}")
        
        except Exception as e:
            logger.error(f"Error processing sheet {sheet_name} for raw data: {str(e)}")
    
    # Now process the raw data to extract structured information
    departments_found = set()  # To keep track of unique departments
    for sheet_data in structured_data["raw_data"]:
        rows = sheet_data["rows"]
        
        # Skip sheets with too few rows
        if len(rows) < 3:
            continue
        
        # Try to identify header rows by looking at rows that might contain column headers
        header_candidates = []
        for i, row in enumerate(rows[:10]):  # Check first 10 rows
            values = list(row.values())
            non_empty = [v for v in values if v and isinstance(v, str) and len(v.strip()) > 0]
            
            # If this looks like a header row (contains keywords like "area", "risk", "control")
            header_keywords = ['area', 'type', 'risk', 'control', 'process', 'objective', 'what can go wrong']
            if any(keyword in v.lower() for v in non_empty for keyword in header_keywords):
                header_candidates.append((i, values))
        
        if header_candidates:
            # Use the last header candidate as our header
            header_idx, header_values = header_candidates[-1]
            
            # Identify column positions
            area_col = None
            control_obj_col = None
            risk_col = None
            area_subprocess_col = None
            
            for i, val in enumerate(header_values):
                val_lower = str(val).lower()
                if any(keyword in val_lower for keyword in ['area', 'department']):
                    area_col = i
                elif any(keyword in val_lower for keyword in ['control objective', 'objective']):
                    control_obj_col = i
                elif any(keyword in val_lower for keyword in ['risk', 'what can go wrong']):
                    risk_col = i
                elif any(keyword in val_lower for keyword in ['sub process', 'subprocess', 'area/']):
                    area_subprocess_col = i
            
            # Process data rows after the header
            for row_idx in range(header_idx + 1, len(rows)):
                row = rows[row_idx]
                values = list(row.values())
                
                # Skip rows that are too short
                if len(values) < max(filter(None, [area_col, control_obj_col, risk_col, area_subprocess_col]), default=0) + 1:
                    continue
                
                # Extract information
                area = values[area_col] if area_col is not None and area_col < len(values) else ""
                control_obj = values[control_obj_col] if control_obj_col is not None and control_obj_col < len(values) else ""
                risk = values[risk_col] if risk_col is not None and risk_col < len(values) else ""
                subprocess = values[area_subprocess_col] if area_subprocess_col is not None and area_subprocess_col < len(values) else ""
                
                # Skip rows without meaningful data
                if not (area or control_obj or risk):
                    continue
                
                # Add department if found
                if area and area.strip():
                    departments_found.add(area.strip())
                
                # Create control objective entry
                if control_obj or risk:
                    # Detect if this is a gap based on risk description
                    is_gap = False
                    risk_text = str(risk).lower()
                    gap_keywords = ['inadequate', 'missing', 'lack', 'absence', 'not adequate', 'incorrect', 
                                   'error', 'without', 'unauthorized', 'risk', 'fail', 'fraud', 'inappropriate']
                    
                    if any(keyword in risk_text for keyword in gap_keywords):
                        is_gap = True
                    
                    # Infer risk level
                    risk_level = "Medium"  # Default
                    high_risk_keywords = ['critical', 'high', 'severe', 'significant', 'major', 'fraud', 'unauthorized', 'incorrect']
                    low_risk_keywords = ['minor', 'low', 'minimal', 'small', 'unlikely']
                    
                    if any(keyword in risk_text for keyword in high_risk_keywords):
                        risk_level = "High"
                    elif any(keyword in risk_text for keyword in low_risk_keywords):
                        risk_level = "Low"
                    
                    # Add control objective
                    objective = {
                        "department": area.strip() if area else "Unknown",
                        "objective": control_obj,
                        "what_can_go_wrong": risk,
                        "risk_level": risk_level,
                        "control_activities": control_obj,
                        "is_gap": is_gap,
                        "gap_details": risk if is_gap else "",
                        "proposed_control": "",
                        "area_subprocess": subprocess
                    }
                    
                    structured_data["control_objectives"].append(objective)
                    
                    # Add gap if applicable
                    if is_gap:
                        gap = {
                            "department": area.strip() if area else "Unknown",
                            "control_objective": control_obj,
                            "gap_title": risk[:50] + "..." if len(risk) > 50 else risk,
                            "description": risk,
                            "risk_impact": risk,
                            "proposed_solution": "",
                            "area_subprocess": subprocess
                        }
                        structured_data["gaps"].append(gap)
        else:
            # If no header candidates, try a more generic approach
            for row in rows:
                values = list(row.values())
                
                # Need at least a few values to be meaningful
                if len([v for v in values if v and v.strip()]) < 3:
                    continue
                
                # Try to find potential department/area names
                for value in values:
                    if not isinstance(value, str) or not value.strip():
                        continue
                        
                    # Check if this looks like a department/area name
                    potential_area = value.strip()
                    area_keywords = ['employee', 'payroll', 'personnel', 'attendance', 'leave', 'management',
                                    'separation', 'maintenance', 'processing', 'department', 'hr', 'finance']
                    
                    if any(keyword in potential_area.lower() for keyword in area_keywords):
                        # This could be a department name
                        if len(potential_area) > 3 and len(potential_area.split()) <= 5:
                            departments_found.add(potential_area)
    
    # Add found departments to structured data
    structured_data["departments"] = list(departments_found)
    
    # If we still don't have enough control objectives, try a more direct approach
    if len(structured_data["control_objectives"]) < 5:
        logger.warning("Few control objectives found, trying direct extraction approach")
        
        # Check all raw data again
        for sheet_data in structured_data["raw_data"]:
            rows = sheet_data["rows"]
            
            for row in rows:
                values = list(row.values())
                
                # Skip rows without enough data
                if len([v for v in values if v and v.strip()]) < 3:
                    continue
                
                # Try to identify columns based on content patterns
                area_val = None
                obj_val = None
                risk_val = None
                
                for val in values:
                    if not isinstance(val, str) or not val.strip():
                        continue
                    
                    val_lower = val.lower()
                    
                    # Employee Master, Payroll, Leave Management likely to be area
                    if any(term in val_lower for term in ['employee master', 'payroll', 'leave management', 'attendance']):
                        area_val = val
                    
                    # Values with "details", "control exists", "review" likely objectives
                    elif any(term in val_lower for term in ['details', 'control', 'review', 'access', 'monitoring']):
                        obj_val = val
                    
                    # Values with negative terms likely risks
                    elif any(term in val_lower for term in ['incorrect', 'unauthorized', 'absence', 'inadequate']):
                        risk_val = val
                
                # If we found area and either objective or risk, create a control objective
                if area_val and (obj_val or risk_val):
                    # Add to departments if needed
                    if area_val not in departments_found:
                        departments_found.add(area_val)
                        structured_data["departments"].append(area_val)
                    
                    # Create control objective
                    objective = {
                        "department": area_val,
                        "objective": obj_val or "Unknown",
                        "what_can_go_wrong": risk_val or "",
                        "risk_level": "Medium",  # Default
                        "control_activities": obj_val or "",
                        "is_gap": bool(risk_val),
                        "gap_details": risk_val or "",
                        "proposed_control": ""
                    }
                    
                    structured_data["control_objectives"].append(objective)
    
    # Ensure we have at least some departments
    if not structured_data["departments"]:
        # Extract from screenshot sample - these are the departments we saw in the image
        default_departments = [
            "Employee Master Maintenance",
            "Attendance & Payroll Processing",
            "Payroll and Personnel",
            "Leave Management",
            "Separation"
        ]
        structured_data["departments"] = default_departments
    
    # Generate summary stats
    structured_data["total_controls"] = len(structured_data["control_objectives"])
    structured_data["control_gaps"] = len(structured_data["gaps"])
    
    # Generate risk distribution
    risk_levels = [obj.get("risk_level", "").strip() for obj in structured_data["control_objectives"] if obj.get("risk_level", "").strip()]
    risk_distribution = {}
    for level in risk_levels:
        if level:
            # Normalize risk levels (High/Medium/Low)
            if level.lower() in ['high', 'h', 'critical', 'severe']:
                std_level = 'High'
            elif level.lower() in ['medium', 'm', 'mod', 'moderate']:
                std_level = 'Medium'
            elif level.lower() in ['low', 'l', 'minor']:
                std_level = 'Low'
            else:
                std_level = 'Medium'  # Default
                
            if std_level in risk_distribution:
                risk_distribution[std_level] += 1
            else:
                risk_distribution[std_level] = 1
    
    # Ensure all three risk levels exist in the distribution
    for level in ['High', 'Medium', 'Low']:
        if level not in risk_distribution:
            risk_distribution[level] = 0
    
    structured_data["risk_distribution"] = risk_distribution
    
    # Specific Risk Types Analysis - based on the user's request
    risk_types = ["Operational", "Financial", "Fraud", "Financial Fraud", "Operational Fraud"]
    risk_type_mapping = {}
    
    # Keywords for each risk type
    risk_type_keywords = {
        "Operational": ["process", "workflow", "efficiency", "performance", "delivery", "resource", "procedure", "operational", "operation"],
        "Financial": ["financial", "budget", "cost", "expense", "revenue", "payment", "accounting", "payroll", "salary"],
        "Fraud": ["fraud", "misappropriation", "theft", "falsification", "bribery", "corruption", "unauthorized"],
        "Financial Fraud": ["financial fraud", "embezzlement", "accounting fraud", "false reporting", "misstatement", "incorrect amount"],
        "Operational Fraud": ["operational fraud", "process manipulation", "override", "unauthorized", "fictitious", "absence of control"]
    }
    
    # Identify risk types for each control objective
    for obj in structured_data["control_objectives"]:
        obj_text = f"{obj.get('objective', '')} {obj.get('what_can_go_wrong', '')}"
        obj_text = obj_text.lower()
        
        # Identify applicable risk types
        obj["risk_types"] = []
        for risk_type, keywords in risk_type_keywords.items():
            if any(keyword in obj_text for keyword in keywords):
                obj["risk_types"].append(risk_type)
                
                # Count risk types
                if risk_type not in risk_type_mapping:
                    risk_type_mapping[risk_type] = 0
                risk_type_mapping[risk_type] += 1
    
    # Add risk type mapping to structured data
    structured_data["risk_type_mapping"] = risk_type_mapping
    
    logger.info(f"Extracted {structured_data['total_controls']} control objectives and {len(structured_data['departments'])} departments")
    logger.info(f"Departments found: {structured_data['departments']}")
    logger.info(f"Identified risk types: {risk_type_mapping}")
    
    return structured_data

def process_csv(file_path: str) -> Dict[str, Any]:
    """Process CSV files and extract RCM data"""
    logger.info(f"Processing CSV file: {file_path}")
    
    # Read CSV
    df = pd.read_csv(file_path)
    
    # Initialize the structured data
    structured_data = {
        "metadata": {
            "file_name": os.path.basename(file_path),
            "file_type": "csv"
        },
        "control_objectives": [],
        "risks": [],
        "controls": [],
        "gaps": [],
        "departments": []
    }
    
    # Clean up column names and drop empty rows
    df.columns = [str(col).strip() for col in df.columns]
    df = df.dropna(how='all').reset_index(drop=True)
    
    # Print column names for debugging
    logger.info(f"Columns in CSV file: {list(df.columns)}")
    
    # Try to identify column names regardless of exact naming
    relevant_column_patterns = {
        'department': ['department', 'dept', 'function', 'area', 'business unit'],
        'control_objective': ['control objective', 'objective', 'control obj'],
        'what_can_go_wrong': ['what can go wrong', 'risk', 'risk description', 'potential risk'],
        'risk_level': ['risk level', 'risk rating', 'risk priority', 'priority', 'severity'],
        'control_activity': ['control activity', 'control', 'mitigating control', 'control description'],
        'control_gap': ['control/design gap', 'gap', 'control gap', 'design gap'],
        'proposed_control': ['proposed control', 'recommendation', 'remediation', 'action plan']
    }
    
    # Map actual column names to standardized names
    column_mapping = {}
    for std_name, patterns in relevant_column_patterns.items():
        for col in df.columns:
            if any(pattern.lower() in col.lower() for pattern in patterns):
                column_mapping[col] = std_name
                break
    
    logger.info(f"Column mapping for CSV: {column_mapping}")
    
    # Check if we've identified enough columns to be an RCM
    if len(column_mapping) >= 3:  # At least 3 relevant columns found
        logger.info(f"CSV appears to be an RCM with {len(column_mapping)} relevant columns")
        
        # Process each row
        for idx, row in df.iterrows():
            # Skip rows that are likely headers or section titles (usually shorter)
            if sum(pd.notna(row)) < 3:
                continue
            
            # Extract data using the column mapping
            obj_data = {}
            for col, std_name in column_mapping.items():
                if pd.notna(row.get(col, None)):
                    obj_data[std_name] = str(row[col])
            
            # If we have enough data to create a control objective
            if 'control_objective' in obj_data or 'what_can_go_wrong' in obj_data:
                # Defaults for required fields
                objective = {
                    "department": obj_data.get('department', 'Unknown'),
                    "objective": obj_data.get('control_objective', 'Unknown'),
                    "what_can_go_wrong": obj_data.get('what_can_go_wrong', ''),
                    "risk_level": obj_data.get('risk_level', 'Medium'),
                    "control_activities": obj_data.get('control_activity', ''),
                    "is_gap": 'control_gap' in obj_data and pd.notna(obj_data.get('control_gap', '')),
                    "gap_details": obj_data.get('control_gap', ''),
                    "proposed_control": obj_data.get('proposed_control', '')
                }
                
                structured_data["control_objectives"].append(objective)
                
                # Add to departments list if not already there
                dept = objective["department"]
                if dept and dept not in structured_data["departments"]:
                    structured_data["departments"].append(dept)
                
                # If there's a gap, add to gaps list
                if objective["is_gap"]:
                    gap = {
                        "department": objective["department"],
                        "control_objective": objective["objective"],
                        "gap_title": objective["gap_details"][:50] + "..." if len(objective["gap_details"]) > 50 else objective["gap_details"],
                        "description": objective["gap_details"],
                        "risk_impact": objective["what_can_go_wrong"],
                        "proposed_solution": objective["proposed_control"]
                    }
                    structured_data["gaps"].append(gap)
    
    # Handle case where no standardized columns were found but there might still be RCM data
    if not structured_data["control_objectives"]:
        logger.warning("No RCM structure found with standard column names. Attempting alternative extraction...")
        
        # If this looks like a header row followed by data rows
        if len(df) >= 2:
            # Use the first row as headers if they weren't already
            potential_headers = df.iloc[0].values
            if any(isinstance(h, str) and len(h) > 3 for h in potential_headers if h is not None):
                # Extract all rows as potential control objectives
                for idx in range(1, len(df)):
                    row = df.iloc[idx]
                    
                    # Skip empty rows
                    if sum(pd.notna(row)) < 3:
                        continue
                    
                    # Create a control objective from whatever data is available
                    objective = {
                        "department": "Unknown",
                        "objective": str(row.iloc[0]) if pd.notna(row.iloc[0]) else "Unknown",
                        "what_can_go_wrong": str(row.iloc[1]) if len(row) > 1 and pd.notna(row.iloc[1]) else "",
                        "risk_level": "Medium",  # Default value
                        "control_activities": str(row.iloc[2]) if len(row) > 2 and pd.notna(row.iloc[2]) else "",
                        "is_gap": len(row) > 3 and pd.notna(row.iloc[3]),
                        "gap_details": str(row.iloc[3]) if len(row) > 3 and pd.notna(row.iloc[3]) else "",
                        "proposed_control": str(row.iloc[4]) if len(row) > 4 and pd.notna(row.iloc[4]) else ""
                    }
                    
                    structured_data["control_objectives"].append(objective)
                    
                    # If we don't have any departments yet, add a default one
                    if not structured_data["departments"]:
                        structured_data["departments"].append("General")
    
    # Generate summary stats
    structured_data["total_controls"] = len(structured_data["control_objectives"])
    structured_data["control_gaps"] = len(structured_data["gaps"])
    
    # Generate risk distribution
    risk_levels = [obj.get("risk_level", "").strip() for obj in structured_data["control_objectives"] if obj.get("risk_level", "").strip()]
    risk_distribution = {}
    for level in risk_levels:
        if level:
            # Normalize risk levels (High/Medium/Low)
            if level.lower() in ['high', 'h', 'critical', 'severe']:
                std_level = 'High'
            elif level.lower() in ['medium', 'm', 'mod', 'moderate']:
                std_level = 'Medium'
            elif level.lower() in ['low', 'l', 'minor']:
                std_level = 'Low'
            else:
                std_level = 'Medium'  # Default
                
            if std_level in risk_distribution:
                risk_distribution[std_level] += 1
            else:
                risk_distribution[std_level] = 1
    
    # Ensure all three risk levels exist in the distribution
    for level in ['High', 'Medium', 'Low']:
        if level not in risk_distribution:
            risk_distribution[level] = 0
    
    structured_data["risk_distribution"] = risk_distribution
    
    # Generate department risk matrix
    if len(structured_data["departments"]) > 0:
        risk_categories = ["Financial", "Operational", "Compliance", "Strategic", "Technological"]
        department_risks = {dept: {cat: 0 for cat in risk_categories} for dept in structured_data["departments"]}
        
        # Populate risk values based on control objectives
        for obj in structured_data["control_objectives"]:
            dept = obj.get("department", "")
            if not dept or dept not in structured_data["departments"]:
                continue
                
            # Set risk level value
            risk_text = obj.get("risk_level", "").lower()
            if risk_text in ['high', 'h', 'critical', 'severe']:
                risk_level_value = 4
            elif risk_text in ['medium', 'm', 'mod', 'moderate']:
                risk_level_value = 3
            else:
                risk_level_value = 2
                
            # Assign to categories based on keywords
            objective_text = obj.get("objective", "").lower()
            what_can_go_wrong = obj.get("what_can_go_wrong", "").lower()
            combined_text = f"{objective_text} {what_can_go_wrong}"
            
            # Financial risks
            if any(term in combined_text for term in ['financ', 'account', 'budget', 'cost', 'expense', 'revenue', 'payment', 'tax', 'audit']):
                department_risks[dept]["Financial"] = max(department_risks[dept]["Financial"], risk_level_value)
                
            # Operational risks
            if any(term in combined_text for term in ['operat', 'process', 'procedur', 'workflow', 'efficien', 'product', 'service', 'delivery']):
                department_risks[dept]["Operational"] = max(department_risks[dept]["Operational"], risk_level_value)
                
            # Compliance risks
            if any(term in combined_text for term in ['comply', 'compliance', 'regulat', 'legal', 'law', 'policy', 'requirement', 'standard']):
                department_risks[dept]["Compliance"] = max(department_risks[dept]["Compliance"], risk_level_value)
                
            # Strategic risks
            if any(term in combined_text for term in ['strateg', 'goal', 'objective', 'mission', 'vision', 'plan', 'market', 'competi']):
                department_risks[dept]["Strategic"] = max(department_risks[dept]["Strategic"], risk_level_value)
                
            # Technological risks
            if any(term in combined_text for term in ['tech', 'system', 'data', 'secur', 'access', 'software', 'hardware', 'it ', 'cyber']):
                department_risks[dept]["Technological"] = max(department_risks[dept]["Technological"], risk_level_value)
        
        # Ensure all departments have at least some risk level for each category
        for dept in structured_data["departments"]:
            for cat in risk_categories:
                if department_risks[dept][cat] == 0:
                    department_risks[dept][cat] = 1  # Set minimum risk level
        
        structured_data["department_risks"] = department_risks
    
    logger.info(f"Extracted {structured_data['total_controls']} control objectives and {len(structured_data['departments'])} departments from CSV")
    
    return structured_data

def process_pdf(file_path: str) -> Dict[str, Any]:
    """Process PDF files and extract RCM data using PDF extraction and LLM later"""
    logger.info(f"Processing PDF file: {file_path}")
    
    try:
        # Extract text from PDF
        extracted_text = ""
        
        # Try with pdfplumber first
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted_text += page.extract_text() + "\n\n"
        
        # If pdfplumber failed or returned empty text, try PyPDF2
        if not extracted_text.strip():
            pdf_reader = PdfReader(file_path)
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n\n"
        
        # Initialize the structured data
        structured_data = {
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_type": "pdf",
                "text_length": len(extracted_text)
            },
            "extracted_text": extracted_text,
            "control_objectives": [],
            "risks": [],
            "controls": [],
            "gaps": [],
            "departments": [],
            "raw_text": True  # Flag to indicate this requires LLM processing
        }
        
        # We'll extract the structured data using LLM later
        # For now, create some placeholder data for sample visualization
        structured_data["total_controls"] = 0
        structured_data["control_gaps"] = 0
        structured_data["risk_score"] = "N/A"
        
        structured_data["risk_distribution"] = {
            "High": 0,
            "Medium": 0, 
            "Low": 0
        }
        
        structured_data["departments"] = ["Finance", "IT", "Operations", "HR"]
        risk_categories = ["Financial", "Operational", "Compliance", "Strategic", "Technological"]
        department_risks = {dept: {cat: 0 for cat in risk_categories} for dept in structured_data["departments"]}
        
        structured_data["department_risks"] = department_risks
        
        return structured_data
        
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        raise

def process_docx(file_path: str) -> Dict[str, Any]:
    """Process DOCX files and extract RCM data"""
    logger.info(f"Processing DOCX file: {file_path}")
    
    try:
        # Extract text from DOCX
        doc = docx.Document(file_path)
        extracted_text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract tables if any
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        # Initialize the structured data
        structured_data = {
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_type": "docx",
                "text_length": len(extracted_text),
                "tables_count": len(tables_data)
            },
            "extracted_text": extracted_text,
            "tables": tables_data,
            "control_objectives": [],
            "risks": [],
            "controls": [],
            "gaps": [],
            "departments": [],
            "raw_text": True  # Flag to indicate this requires LLM processing
        }
        
        # We'll extract the structured data using LLM later
        # For now, create some placeholder data for sample visualization
        structured_data["total_controls"] = 0
        structured_data["control_gaps"] = 0
        structured_data["risk_score"] = "N/A"
        
        structured_data["risk_distribution"] = {
            "High": 0,
            "Medium": 0, 
            "Low": 0
        }
        
        structured_data["departments"] = ["Finance", "IT", "Operations", "HR"]
        risk_categories = ["Financial", "Operational", "Compliance", "Strategic", "Technological"]
        department_risks = {dept: {cat: 0 for cat in risk_categories} for dept in structured_data["departments"]}
        
        structured_data["department_risks"] = department_risks
        
        return structured_data
        
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
        raise 